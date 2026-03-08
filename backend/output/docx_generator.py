"""
Intelli-Credit — python-docx CAM Generator

Generates a Credit Appraisal Memorandum in .docx format following
Indian banking standard layout:
  - Header block (company info)
  - Executive Summary with score
  - Hard Blocks (if any)
  - Score Breakdown by Module (6 modules)
  - Detailed Metric Analysis with source tracing
  - Loan Terms Recommendation
  - Cross-Verification Summary
  - Risk Flags
  - Tickets
  - Footer

The generator accepts a context dict (same shape as Jinja2 template context)
and produces a .docx file at the specified output path.
"""

import io
import logging
import os
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────
# Style Constants
# ──────────────────────────────────────────────

_TEAL = RGBColor(0x0D, 0x94, 0x88)      # teal-600
_RED = RGBColor(0xDC, 0x26, 0x26)        # red-600
_GREEN = RGBColor(0x05, 0x96, 0x69)      # emerald-600
_AMBER = RGBColor(0xD9, 0x77, 0x06)      # amber-600
_GRAY = RGBColor(0x64, 0x74, 0x8B)       # slate-500
_BLACK = RGBColor(0x1E, 0x29, 0x3B)      # slate-800
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _score_color(score: int) -> RGBColor:
    if score >= 750:
        return _GREEN
    elif score >= 650:
        return _TEAL
    elif score >= 550:
        return _AMBER
    elif score >= 350:
        return _RED
    return RGBColor(0x99, 0x1B, 0x1B)


def _impact_color(impact: int) -> RGBColor:
    return _GREEN if impact > 0 else _RED


# ──────────────────────────────────────────────
# Helper Functions
# ──────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _BLACK
    return h


def _add_table(doc: Document, headers: List[str], rows: List[List[str]],
               col_widths: Optional[List[float]] = None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            row.cells[c_idx].text = str(text)
            for paragraph in row.cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def _add_key_value(doc: Document, key: str, value: str):
    """Add a bold key: value paragraph."""
    p = doc.add_paragraph()
    run_key = p.add_run(f"{key}: ")
    run_key.bold = True
    run_key.font.size = Pt(10)
    run_val = p.add_run(str(value))
    run_val.font.size = Pt(10)
    return p


# ──────────────────────────────────────────────
# Main Generator
# ──────────────────────────────────────────────

def generate_cam_docx(context: Dict[str, Any], output_path: Optional[str] = None) -> bytes:
    """Generate a CAM .docx document.

    Args:
        context: Dict with keys: company_name, sector, loan_type, loan_amount,
                 session_id, date, score, score_band, outcome, recommendation,
                 base_score, modules (list of ScoreModuleSummary dicts),
                 hard_blocks, loan_terms, cross_verifications, risk_flags, tickets
        output_path: Optional file path to save. If None, returns bytes only.

    Returns:
        The .docx content as bytes.
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Title ──
    title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = _BLACK

    subtitle = doc.add_paragraph("Confidential — For Internal Use Only")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = _GRAY
        run.italic = True

    doc.add_paragraph()  # spacer

    # ── Header Info Table ──
    company_name = context.get("company_name", "Unknown")
    header_data = [
        ["Company", company_name],
        ["Sector", context.get("sector", "N/A")],
        ["Facility", context.get("loan_type", "N/A")],
        ["Amount", context.get("loan_amount", "N/A")],
        ["Assessment ID", context.get("session_id", "N/A")],
        ["Date", context.get("date", datetime.utcnow().strftime("%Y-%m-%d"))],
    ]
    _add_table(doc, ["Field", "Value"], header_data, col_widths=[5, 12])
    doc.add_paragraph()

    score = context.get("score", 0)
    score_band = context.get("score_band", "N/A")
    outcome = context.get("outcome", "PENDING")

    # ── Section 1: Executive Summary ──
    _add_heading(doc, "1. Executive Summary", level=1)
    p_score = doc.add_paragraph()
    p_score.add_run("Credit Score: ").bold = True
    run_score = p_score.add_run(f"{score} / 850")
    run_score.font.size = Pt(14)
    run_score.font.color.rgb = _score_color(score)
    run_score.bold = True

    _add_key_value(doc, "Score Band", score_band)
    _add_key_value(doc, "Outcome", outcome)

    recommendation = context.get("recommendation", "")
    if recommendation:
        _add_key_value(doc, "Recommendation", recommendation)

    # ── Hard Blocks (if any) ──
    hard_blocks = context.get("hard_blocks", [])
    if hard_blocks:
        _add_heading(doc, "⚠ Hard Block Triggers", level=2)
        rows = []
        for hb in hard_blocks:
            rows.append([
                _get(hb, "trigger"),
                str(_get(hb, "score_cap")),
                _get(hb, "evidence"),
                _get(hb, "source"),
            ])
        _add_table(doc, ["Trigger", "Score Cap", "Evidence", "Source"], rows)
        doc.add_paragraph()

    # ── Section 2: Score Breakdown by Module ──
    modules = context.get("modules", [])
    _add_heading(doc, "2. Score Breakdown", level=1)
    if modules:
        mod_rows = []
        for mod in modules:
            mod_rows.append([
                _get(mod, "module"),
                str(_get(mod, "score", 0)),
                f"+{_get(mod, 'max_positive', 0)}",
                str(_get(mod, "max_negative", 0)),
                str(len(_get(mod, "metrics", []))),
            ])
        _add_table(doc, ["Module", "Score", "Max (+)", "Max (−)", "Metrics"], mod_rows)

    base_score = context.get("base_score", 350)
    _add_key_value(doc, "Base Score", str(base_score))
    _add_key_value(doc, "Final Score", f"{score} / 850")
    doc.add_paragraph()

    # ── Section 3: Detailed Metric Analysis ──
    _add_heading(doc, "3. Detailed Metric Analysis", level=1)
    for mod in modules:
        mod_name = _get(mod, "module", "Unknown")
        mod_score = _get(mod, "score", 0)
        _add_heading(doc, f"{mod_name} ({mod_score} pts)", level=2)

        metrics = _get(mod, "metrics", [])
        if metrics:
            metric_rows = []
            for m in metrics:
                impact = _get(m, "score_impact", 0)
                sign = "+" if impact > 0 else ""
                confidence = _get(m, "confidence", 0)
                conf_pct = f"{confidence * 100:.0f}%" if isinstance(confidence, (int, float)) else str(confidence)
                metric_rows.append([
                    _get(m, "metric_name"),
                    _get(m, "metric_value"),
                    f"{sign}{impact}",
                    f"{_get(m, 'source_document')} p.{_get(m, 'source_page')}",
                    conf_pct,
                    _get(m, "reasoning"),
                ])
            _add_table(
                doc,
                ["Metric", "Value", "Impact", "Source", "Conf.", "Reasoning"],
                metric_rows,
            )
        doc.add_paragraph()

    # ── Section 4: Loan Terms ──
    loan_terms = context.get("loan_terms")
    if loan_terms:
        _add_heading(doc, "4. Recommended Loan Terms", level=1)
        lt_rows = [
            ["Sanction %", f"{_get(loan_terms, 'sanction_pct')}%"],
            ["Rate", _get(loan_terms, "rate")],
            ["Tenure", _get(loan_terms, "tenure")],
            ["Review Frequency", _get(loan_terms, "review")],
        ]
        _add_table(doc, ["Parameter", "Value"], lt_rows, col_widths=[5, 12])
        doc.add_paragraph()

    # ── Section 5: Cross-Verification Summary ──
    cross_verifications = context.get("cross_verifications", [])
    if cross_verifications:
        _add_heading(doc, "5. Cross-Verification Summary", level=1)
        cv_rows = []
        for cv in cross_verifications:
            cv_rows.append([
                _get(cv, "check"),
                _get(cv, "sources"),
                _get(cv, "result"),
                _get(cv, "discrepancy"),
            ])
        _add_table(doc, ["Check", "Sources", "Result", "Discrepancy"], cv_rows)
        doc.add_paragraph()

    # ── Section 6: Risk Flags ──
    risk_flags = context.get("risk_flags", [])
    if risk_flags:
        _add_heading(doc, "6. Risk Flags", level=1)
        rf_rows = []
        for rf in risk_flags:
            impact = _get(rf, "score_impact", 0)
            severity = "Critical" if impact < -20 else "Warning"
            rf_rows.append([
                _get(rf, "metric_name"),
                severity,
                str(impact),
                f"{_get(rf, 'source_document')} p.{_get(rf, 'source_page')}",
            ])
        _add_table(doc, ["Flag", "Severity", "Impact", "Source"], rf_rows)
        doc.add_paragraph()

    # ── Section 7: Tickets ──
    tickets = context.get("tickets", [])
    if tickets:
        _add_heading(doc, "7. Tickets Raised", level=1)
        t_rows = []
        for t in tickets:
            t_rows.append([
                _get(t, "title"),
                _get(t, "severity"),
                _get(t, "status"),
                str(_get(t, "score_impact", 0)),
            ])
        _add_table(doc, ["Title", "Severity", "Status", "Score Impact"], t_rows)
        doc.add_paragraph()

    # ── Footer ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer.add_run(
        f"Generated by Intelli-Credit AI Engine | "
        f"{context.get('date', datetime.utcnow().strftime('%Y-%m-%d'))}"
    )
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = _GRAY
    run_f.italic = True

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = disclaimer.add_run(
        "This document is machine-generated and subject to credit officer review."
    )
    run_d.font.size = Pt(8)
    run_d.font.color.rgb = _GRAY
    run_d.italic = True

    # ── Save / Return ──
    buffer = io.BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(content)
        logger.info("[CAM-DOCX] Saved to %s (%d bytes)", output_path, len(content))

    return content


def _get(obj, key, default="N/A"):
    """Safely get a value from dict or object."""
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


# ──────────────────────────────────────────────
# Convenience: Build context from ScoreResponse
# ──────────────────────────────────────────────

def build_cam_context(score_response, assessment=None) -> Dict[str, Any]:
    """Build a CAM template context from a ScoreResponse and optional AssessmentSummary.

    Works with both Pydantic models and plain dicts.
    """
    sr = score_response
    ctx = {
        "company_name": _get(sr, "company_name", "Unknown"),
        "session_id": _get(sr, "session_id", ""),
        "score": _get(sr, "score", 0),
        "score_band": str(_get(sr, "score_band", "N/A")),
        "outcome": str(_get(sr, "outcome", "PENDING")),
        "recommendation": _get(sr, "recommendation", ""),
        "base_score": _get(sr, "base_score", 350),
        "date": datetime.utcnow().strftime("%Y-%m-%d"),
    }

    # Modules — convert Pydantic models to dicts for template
    modules = _get(sr, "modules", [])
    ctx["modules"] = [
        m.model_dump() if hasattr(m, "model_dump") else m
        for m in modules
    ]

    # Hard blocks
    hb = _get(sr, "hard_blocks", [])
    ctx["hard_blocks"] = [
        h.model_dump() if hasattr(h, "model_dump") else h
        for h in hb
    ]

    # Loan terms
    lt = _get(sr, "loan_terms", None)
    if lt:
        ctx["loan_terms"] = lt.model_dump() if hasattr(lt, "model_dump") else lt

    # From assessment (optional extras)
    if assessment:
        ctx["sector"] = _get(_get(assessment, "company", {}), "sector", "N/A")
        ctx["loan_type"] = _get(_get(assessment, "company", {}), "loan_type", "N/A")
        ctx["loan_amount"] = _get(_get(assessment, "company", {}), "loan_amount", "N/A")
        ctx["tickets"] = [
            t.model_dump() if hasattr(t, "model_dump") else t
            for t in _get(assessment, "tickets", [])
        ]

    # Risk flags — metrics with score_impact < -10
    risk_flags = []
    for mod in ctx.get("modules", []):
        for m in _get(mod, "metrics", []):
            if _get(m, "score_impact", 0) < -10:
                risk_flags.append(m)
    ctx["risk_flags"] = risk_flags

    return ctx
